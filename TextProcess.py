# 用于将文本文件中的语句分段读入mysql数据库
import docx
import re

import pymysql


class TextProcess:
    def __init__(self):
        self.filePath = "D:\\知识图谱录入文件\\计算机组成与体系结构.docx"
        self.password = "lzg19981202"
        self.host = "localhost"
        self.user = "root"
        self.db = "knowledge_graph"
        self.tableName = "paragraphs"

        # 非标题段落最小长度
        self.min_len = 25
        # 将要读入的所有数据
        self.items = []

    def ReadFile(self):
        # 记录段落的标题
        titles = {}
        # 记录段落的位置，根据标题前面的序号记录
        index = ""
        # 每个要存入的数据对象

        file = docx.Document(self.filePath)
        print("当前该上传文件中共有段落数：" + str(len(file.paragraphs)))

        for i in range(len(file.paragraphs)):
            # item要放入循环，不然每次都会修改之前的
            item = {}
            if file.paragraphs[i].text.isspace() or len(file.paragraphs[i].text) == 0:
                print("当前段落为空，跳过不读入")
                continue
            if file.paragraphs[i].text[0].isdigit() and "." in file.paragraphs[i].text:
                print("当前段落识别为标题")
                index = re.sub('[\u4e00-\u9fa5]', '', file.paragraphs[i].text)
                title = file.paragraphs[i].text.split(".")
                if len(title) == 2:
                    print("一级标题")
                    titles["一级标题"] = title[-1].strip()
                elif len(title) == 3:
                    print("二级标题")
                    titles["二级标题"] = title[-1].strip()
                elif len(title) == 4:
                    print("三级标题")
                    titles["三级标题"] = title[-1].strip()
                elif len(title) == 5:
                    print("四级标题")
                    titles["四级标题"] = title[-1].strip()
            if len(file.paragraphs[i].text) < self.min_len:
                print("读取到的段落太短，不进行实体识别")
                continue
            item["file_name"] = "计算机组成与体系结构.docx"
            item["title"] = "--".join(titles.values())
            item["position"] = index
            item["content"] = file.paragraphs[i].text
            self.items.append(item)

    def insertDB(self):
        print("开始插入数据库")
        print(self.items)
        db = pymysql.connect(host=self.host, user=self.user, password=self.password, db=self.db)
        cursor = db.cursor()
        for item in self.items:
            sql = f"INSERT INTO {self.tableName} (file_name, title, content, position) VALUES ('{item['file_name']}', '{item['title']}', '{item['content']}', '{item['position']}')"
            print(sql)
            try:
                cursor.execute(sql)
                db.commit()
            except:
                print("插入失败")
                db.rollback()
        cursor.close()
        db.close()


if __name__ == '__main__':
    r = TextProcess()
    r.ReadFile()
    r.insertDB()
